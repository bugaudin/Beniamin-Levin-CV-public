#!/usr/bin/env python3
"""Generate cover letter PDF from plain text via LibreOffice."""
import argparse
import os
import re
import shutil
import subprocess
import tempfile

from docx import Document
from docx.shared import Pt, Inches

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SOFFICE = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
FONT_BODY = 'Avenir Next'
FONT_SIZE = 11.0


def convert_docx_to_pdf(docx_path, pdf_path):
    outdir = os.path.dirname(pdf_path)
    subprocess.run(
        [SOFFICE, '--headless', '--convert-to', 'pdf', '--outdir', outdir, docx_path],
        check=True,
        capture_output=True,
    )
    generated = os.path.join(
        outdir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
    )
    if generated != pdf_path:
        shutil.move(generated, pdf_path)


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    lines = text.split('\n')
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = FONT_BODY
        run.font.size = Pt(FONT_SIZE)
        if index < len(lines) - 1:
            run.add_break()


def build_cover_letter_pdf(input_txt, output_pdf):
    with open(input_txt, encoding='utf-8') as f:
        text = f.read().strip()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    paragraphs = re.split(r'\n\s*\n', text)
    for paragraph in paragraphs:
        stripped = paragraph.strip()
        if stripped:
            add_paragraph(doc, stripped)

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, 'cover-letter.docx')
        doc.save(docx_path)
        convert_docx_to_pdf(docx_path, output_pdf)

    print(f'saved {output_pdf}')


def main():
    parser = argparse.ArgumentParser(description='Generate cover letter PDF from plain text')
    parser.add_argument('--input', required=True, help='Source plain-text cover letter')
    parser.add_argument('--output', required=True, help='Output PDF path')
    args = parser.parse_args()

    input_path = os.path.join(SCRIPT_DIR, args.input)
    output_path = os.path.join(SCRIPT_DIR, args.output)
    build_cover_letter_pdf(input_path, output_path)


if __name__ == '__main__':
    main()
