#!/usr/bin/env python3
"""Generate styled A4 CV PDF from README.md via LibreOffice."""
import argparse
import os
import re
import shutil
import subprocess
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1A, 0x3A, 0x52)
ACCENT_HEX = '1A3A52'
GREY = RGBColor(0x5C, 0x63, 0x6A)
BLACK = RGBColor(0x1F, 0x24, 0x28)
PAGE_W_IN = 8.27
MARGIN_LR = Inches(0.38)
MARGIN_TOP = Inches(0.22)
MARGIN_BOTTOM = Inches(0.26)
LINE_SPACING = 1.02

# Distinctive Mac-native fonts that LibreOffice can embed.
FONT_BODY = 'Avenir Next'
FONT_DISPLAY = 'Georgia'

# Base sizes at scale 1.0; scaled at build time for page-break tuning.
BASE_FONT_BODY = 9.4
BASE_FONT_SECTION = 10.5
BASE_FONT_JOB = 10.0
BASE_FONT_JOB_DATE = 9.2
BASE_FONT_ROLE = 9.3
BASE_FONT_SUBHEAD = 9.5
BASE_FONT_NAME = 22.0
BASE_FONT_CONTACT = 8.8
BASE_FONT_EDU = 9.5
BASE_FONT_EDU_DETAIL = 9.1

SECTION_TITLES = {
    'PROFESSIONAL SUMMARY',
    'CORE COMPETENCIES & TECHNICAL SKILLS',
    'CORE TECHNICAL SKILLS',
    'PROFESSIONAL EXPERIENCE',
    'FREELANCE & PERSONAL PROJECTS — PART-TIME / OUTSIDE FULL-TIME EMPLOYMENT',
    'EDUCATION & ACADEMIC TRAINING',
    'EDUCATION',
    'LANGUAGES',
}

SKILLS_SECTIONS = {
    'CORE COMPETENCIES & TECHNICAL SKILLS',
    'CORE TECHNICAL SKILLS',
}
EDUCATION_SECTIONS = {
    'EDUCATION & ACADEMIC TRAINING',
    'EDUCATION',
}

FREELANCE_SECTION = 'Freelance & Personal Projects — Part-Time / Outside Full-Time Employment'
FREELANCE_MARKER = 'freelance'
RELOCATION_LINE = (
    'Seeking relocation to the Netherlands for a Tech Lead role with an '
    'IND-recognised employer under Highly Skilled Migrant sponsorship.'
)
DEFAULT_OUTPUT = 'Beniamin-Levin-CV.pdf'
README_NAME = 'README-OLD.md'
HANDS_ON_README = 'README.md'
HANDS_ON_OUTPUT = 'Beniamin-Levin-CV-hands-on.pdf'
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SOFFICE = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
INLINE_MARKUP_RE = re.compile(
    r'\[([^\]]+)\]\(([^)]+)\)|\*\*(.+?)\*\*|\bikite\.fyi\b|https?://[^\s),;]+'
)


def tokenize_inline_markup(text):
    parts = []
    pos = 0
    for match in INLINE_MARKUP_RE.finditer(text):
        if match.start() > pos:
            parts.append(('text', text[pos:match.start()], False))
        token = match.group(0)
        if token.startswith('['):
            url = match.group(2)
            parts.append(('link', url, url, False))
        elif token.startswith('**'):
            parts.append(('text', match.group(3), True))
        elif token == 'ikite.fyi':
            parts.append(('link', 'ikite.fyi', 'https://ikite.fyi', False))
        elif token.startswith('http'):
            url = token.rstrip('.,;')
            parts.append(('link', url, url, False))
            if len(url) < len(token):
                parts.append(('text', token[len(url):], False))
        pos = match.end()
    if pos < len(text):
        parts.append(('text', text[pos:], False))
    return parts


def load_secrets():
    secrets_path = os.path.join(SCRIPT_DIR, 'secrets.txt')
    secrets = {}
    try:
        with open(secrets_path, encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                key, _, value = line.partition('=')
                key = key.strip()
                value = value.strip()
                if key:
                    secrets[key] = value
    except FileNotFoundError:
        pass
    if not secrets:
        raise SystemExit('secrets.txt not found or empty — add key=value entries (e.g. PHONE=...)')
    return secrets


def strip_md_bold(text):
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text).strip()


def strip_md_italic(text):
    text = text.strip()
    if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
        return text[1:-1].strip()
    return strip_md_bold(text)


def is_section_header(text):
    return strip_md_bold(text).upper() in SECTION_TITLES


def is_italic_line(text):
    return (
        text.startswith('*')
        and text.endswith('*')
        and not text.startswith('**')
        and not text.startswith('* ')
    )


BULLET_PREFIXES = ('* ', '- ')


def is_bullet_line(text):
    return text.startswith(BULLET_PREFIXES)


def strip_bullet_prefix(text):
    text = text.strip()
    for prefix in BULLET_PREFIXES:
        if text.startswith(prefix):
            return text[len(prefix):]
    return text


def parse_bullet_segments(text):
    text = strip_bullet_prefix(text)
    match = re.match(r'\*\*(.+?):\*\*\s*(.*)', text, re.DOTALL)
    if match:
        return [(match.group(1) + ': ', True), (match.group(2), False)]
    return [(text, False)]


def parse_skill_line(text):
    match = re.match(r'\*\*(.+?):\*\*\s*(.*)', text.strip())
    if not match:
        return None
    return match.group(1), match.group(2)


def parse_job_header(text):
    inner = strip_md_bold(text)
    if '\t' in inner:
        title, dates = inner.split('\t', 1)
        return title.strip(), dates.strip()
    if ' | ' in inner:
        title, dates = inner.rsplit(' | ', 1)
        return title.strip(), dates.strip()
    return inner, None


def heading_level(text):
    match = re.match(r'^(#{1,6})\s+(.+)$', text)
    if not match:
        return None, None
    return len(match.group(1)), match.group(2).strip()


def extract_markdown_link(text):
    match = re.search(r'\[([^\]]+)\]\(([^)]+)\)', text)
    if not match:
        return None
    return match.group(1), match.group(2)


def is_section_title_text(text):
    return strip_md_bold(text).upper() in SECTION_TITLES


def parse_role_date_line(text):
    """Parse '**Role** | dates' or '*Role*' lines."""
    stripped = text.strip()
    if ' | ' in stripped:
        role_part, dates = stripped.rsplit(' | ', 1)
        return strip_md_bold(role_part), dates.strip()
    if is_italic_line(stripped):
        return strip_md_italic(stripped), None
    if stripped.startswith('**') and stripped.endswith('**'):
        return strip_md_bold(stripped), None
    return None, None


def parse_readme(path):
    with open(path, encoding='utf-8') as f:
        lines = [line.rstrip() for line in f]

    data = {
        'name': None,
        'tagline': None,
        'contact_bits': [],
        'linkedin': None,
        'sections': [],
    }
    current = None

    def start_section(title):
        nonlocal current
        current = {'title': strip_md_bold(title), 'blocks': []}
        data['sections'].append(current)

    def current_title_upper():
        return current['title'].upper() if current else ''

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        level, heading = heading_level(stripped)

        # ATX name / sections / jobs
        if level == 1 and data['name'] is None:
            data['name'] = strip_md_bold(heading)
            i += 1
            continue

        if level == 2 and is_section_title_text(heading):
            start_section(heading)
            i += 1
            continue

        if level == 3 and current_title_upper() == 'PROFESSIONAL EXPERIENCE':
            company = strip_md_bold(heading)
            role = None
            dates = None
            bullets = []
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    j += 1
                    continue
                nxt_level, nxt_heading = heading_level(nxt)
                if nxt_level in {2, 3}:
                    break
                if is_section_title_text(nxt):
                    break
                if role is None and (
                    nxt.startswith('**') or is_italic_line(nxt)
                ):
                    role, dates = parse_role_date_line(nxt)
                    j += 1
                    continue
                if is_bullet_line(nxt):
                    bullets.append(parse_bullet_segments(nxt))
                    j += 1
                    continue
                break
            current['blocks'].append({
                'type': 'job',
                'company': company,
                'dates': dates,
                'role': role,
                'subtitle': None,
                'bullets': bullets,
            })
            i = j
            continue

        # Legacy **NAME** header
        if stripped.startswith('**') and stripped.endswith('**') and data['name'] is None:
            data['name'] = strip_md_bold(stripped)
            i += 1
            continue

        # Tagline under name (ATX format)
        if (
            data['name']
            and data['tagline'] is None
            and not data['contact_bits']
            and not data['sections']
            and stripped.startswith('**')
            and stripped.endswith('**')
            and not is_section_title_text(stripped)
        ):
            data['tagline'] = strip_md_bold(stripped)
            i += 1
            continue

        # Contact line(s)
        if data['name'] and not data['contact_bits'] and not data['sections']:
            link = extract_markdown_link(stripped)
            if link and stripped.startswith('[') and '](' in stripped and '|' not in stripped:
                data['linkedin'] = link
                i += 1
                continue
            if '@' in stripped or '|' in stripped or link:
                if link and data['linkedin'] is None:
                    data['linkedin'] = link
                data['contact_bits'].append(stripped)
                i += 1
                continue

        if stripped.startswith('[') and '](' in stripped and data['linkedin'] is None:
            link = extract_markdown_link(stripped)
            if link:
                data['linkedin'] = link
            i += 1
            continue

        skill = parse_skill_line(strip_bullet_prefix(stripped))
        if skill and current and current_title_upper() in SKILLS_SECTIONS:
            current['blocks'].append({'type': 'skill', 'label': skill[0], 'rest': skill[1]})
            i += 1
            continue

        if stripped.startswith('**'):
            title = strip_md_bold(stripped)
            if stripped.endswith('**') and is_section_title_text(stripped):
                start_section(title)
                i += 1
                continue

            if current and current_title_upper() in {
                'PROFESSIONAL EXPERIENCE',
                FREELANCE_SECTION.upper(),
            } and (stripped.endswith('**') or '\t' in stripped):
                company, dates = parse_job_header(stripped)
                role = None
                subtitle = None
                bullets = []
                j = i + 1
                while j < len(lines):
                    nxt = lines[j].strip()
                    if not nxt:
                        j += 1
                        continue
                    if is_section_title_text(nxt) or heading_level(nxt)[0]:
                        break
                    if is_bullet_line(nxt) and not nxt.startswith('**'):
                        bullets.append(parse_bullet_segments(nxt))
                        j += 1
                        continue
                    if is_italic_line(nxt):
                        if role is None:
                            role = strip_md_italic(nxt)
                        else:
                            subtitle = strip_md_italic(nxt)
                        j += 1
                        continue
                    if nxt.startswith('**'):
                        break
                    break
                current['blocks'].append({
                    'type': 'job',
                    'company': company,
                    'dates': dates,
                    'role': role,
                    'subtitle': subtitle,
                    'bullets': bullets,
                })
                i = j
                continue

            if current and current_title_upper() in EDUCATION_SECTIONS:
                company, dates = parse_job_header(stripped)
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                next_line = lines[j].strip() if j < len(lines) else ''
                role = None
                if (
                    next_line
                    and not is_section_title_text(next_line)
                    and not heading_level(next_line)[0]
                    and not (
                        next_line.startswith('**')
                        and ('|' in next_line or next_line.endswith('**'))
                    )
                ):
                    if is_italic_line(next_line):
                        role = strip_md_italic(next_line)
                    else:
                        role = next_line
                    j += 1
                current['blocks'].append({
                    'type': 'job',
                    'company': company,
                    'dates': dates,
                    'role': role,
                    'subtitle': None,
                    'bullets': [],
                })
                i = j
                continue

        if is_bullet_line(stripped):
            if current:
                current['blocks'].append({
                    'type': 'bullet',
                    'segments': parse_bullet_segments(stripped),
                })
            i += 1
            continue

        if current and current_title_upper() == 'PROFESSIONAL SUMMARY':
            current['blocks'].append({'type': 'para', 'text': stripped})
            i += 1
            continue

        if current and current_title_upper() == 'LANGUAGES':
            current['blocks'].append({'type': 'para', 'text': stripped})
            i += 1
            continue

        i += 1

    return data


def filter_cv_data(data, exclude_paragraphs=(), exclude_sections=()):
    if not exclude_paragraphs and not exclude_sections:
        return data
    exclude_sections_upper = {title.upper() for title in exclude_sections}
    filtered = {
        'name': data['name'],
        'tagline': data.get('tagline'),
        'contact_bits': list(data['contact_bits']),
        'linkedin': data['linkedin'],
        'sections': [],
    }
    for section in data['sections']:
        if section['title'].upper() in exclude_sections_upper:
            continue
        blocks = []
        for block in section['blocks']:
            if block.get('type') == 'para' and block.get('text') in exclude_paragraphs:
                continue
            blocks.append(block)
        filtered['sections'].append({'title': section['title'], 'blocks': blocks})
    return filtered


def scaled_pt(base, scale):
    return Pt(base * scale)


def style_run(run, font_name, size=None, bold=None, italic=None, color=None, tracking=None):
    """Apply font metrics in a LibreOffice-friendly way."""
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:cs'), font_name)
    r_fonts.set(qn('w:eastAsia'), font_name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if tracking is not None:
        spacing = r_pr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            r_pr.append(spacing)
        spacing.set(qn('w:val'), str(tracking))
    return run


class CvBuilder:
    def __init__(self, scale, line_spacing=LINE_SPACING):
        self.scale = scale
        self.line_spacing = line_spacing
        self.font_body = scaled_pt(BASE_FONT_BODY, scale)
        self.font_section = scaled_pt(BASE_FONT_SECTION, scale)
        self.font_job = scaled_pt(BASE_FONT_JOB, scale)
        self.font_job_date = scaled_pt(BASE_FONT_JOB_DATE, scale)
        self.font_role = scaled_pt(BASE_FONT_ROLE, scale)
        self.font_subhead = scaled_pt(BASE_FONT_SUBHEAD, scale)
        self.font_name = scaled_pt(BASE_FONT_NAME, scale)
        self.font_contact = scaled_pt(BASE_FONT_CONTACT, scale)
        self.font_edu = scaled_pt(BASE_FONT_EDU, scale)
        self.font_edu_detail = scaled_pt(BASE_FONT_EDU_DETAIL, scale)
        self.content_w_in = PAGE_W_IN - 2 * MARGIN_LR.inches
        self.doc = Document()
        self._init_styles()

    def _init_styles(self):
        normal = self.doc.styles['Normal']
        normal.font.name = FONT_BODY
        normal.font.size = self.font_body
        normal.font.color.rgb = BLACK
        r_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn('w:ascii'), FONT_BODY)
        r_fonts.set(qn('w:hAnsi'), FONT_BODY)
        r_fonts.set(qn('w:cs'), FONT_BODY)
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = self.line_spacing

        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.top_margin = MARGIN_TOP
        sec.bottom_margin = MARGIN_BOTTOM
        sec.left_margin = MARGIN_LR
        sec.right_margin = MARGIN_LR

    def set_space(self, p, before=0, after=0):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        return p

    def justify(self, p):
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def add_bottom_border(self, p):
        ppr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), ACCENT_HEX)
        pbdr.append(bottom)
        ppr.append(pbdr)

    def name_heading(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_space(p, 0, 3)
        r = p.add_run(text.title())
        style_run(
            r,
            FONT_DISPLAY,
            size=self.font_name,
            bold=True,
            color=ACCENT,
            tracking=40,
        )

    def tagline_line(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_space(p, 0, 3)
        r = p.add_run(text)
        style_run(r, FONT_BODY, size=self.font_role, bold=True, color=BLACK)

    def contact_line(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_space(p, 0, 2)
        r = p.add_run(text)
        style_run(r, FONT_BODY, size=self.font_contact, color=GREY)

    def append_hyperlink(
        self, paragraph, display, url, font_size, bold=False, color=GREY, font_name=FONT_BODY
    ):
        r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        r_fonts = OxmlElement('w:rFonts')
        r_fonts.set(qn('w:ascii'), font_name)
        r_fonts.set(qn('w:hAnsi'), font_name)
        r_fonts.set(qn('w:cs'), font_name)
        r_pr.append(r_fonts)
        if bold:
            r_pr.append(OxmlElement('w:b'))
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), str(color))
        r_pr.append(color_el)
        size = OxmlElement('w:sz')
        size.set(qn('w:val'), str(int(font_size.pt * 2)))
        r_pr.append(size)
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        r_pr.append(underline)
        run.append(r_pr)
        text_el = OxmlElement('w:t')
        text_el.text = display
        run.append(text_el)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def append_rich_text(
        self, paragraph, text, font_size, bold=False, color=BLACK, font_name=FONT_BODY
    ):
        for kind, *payload in tokenize_inline_markup(text):
            if kind == 'text':
                content, span_bold = payload
                r = paragraph.add_run(content)
                style_run(
                    r, font_name, size=font_size, bold=bold or span_bold, color=color
                )
            else:
                display, url, span_bold = payload
                self.append_hyperlink(
                    paragraph,
                    display,
                    url,
                    font_size,
                    bold=bold or span_bold,
                    color=color,
                    font_name=font_name,
                )

    def contact_link(self, text, url):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_space(p, 0, 2)
        self.append_hyperlink(p, text, url, self.font_contact, bold=False, color=GREY)

    def section(self, title, page_break_before=False):
        p = self.doc.add_paragraph()
        if page_break_before:
            p.paragraph_format.page_break_before = True
        self.set_space(p, 7, 3)
        r = p.add_run(title.upper())
        style_run(
            r,
            FONT_BODY,
            size=self.font_section,
            bold=True,
            color=ACCENT,
            tracking=40,
        )
        self.add_bottom_border(p)

    def skill_line(self, label, rest):
        p = self.justify(self.doc.add_paragraph())
        self.set_space(p, 0.9, 0.9)
        r = p.add_run(label + ': ')
        style_run(r, FONT_BODY, size=self.font_body, bold=True, color=ACCENT)
        r2 = p.add_run(rest)
        style_run(r2, FONT_BODY, size=self.font_body, color=BLACK)

    def dated_line(self, title, dates, title_font=None, date_bold=True, page_break_before=False):
        if title_font is None:
            title_font = self.font_job
        p = self.doc.add_paragraph()
        if page_break_before:
            p.paragraph_format.page_break_before = True
        self.set_space(p, 3.5, 0)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(self.content_w_in), WD_TAB_ALIGNMENT.RIGHT
        )
        self.append_rich_text(p, title, title_font, bold=True, color=BLACK)
        if dates:
            rd = p.add_run('\t' + dates)
            style_run(
                rd,
                FONT_BODY,
                size=self.font_job_date,
                bold=date_bold,
                color=GREY,
            )

    def job(self, company, role, dates, subtitle=None, page_break_before=False):
        self.dated_line(
            company, dates, page_break_before=page_break_before
        )
        if role:
            pr = self.doc.add_paragraph()
            self.set_space(pr, 0, 1.2)
            ri = pr.add_run(role)
            style_run(ri, FONT_BODY, size=self.font_role, italic=True, color=ACCENT)
        if subtitle:
            ps = self.doc.add_paragraph()
            self.set_space(ps, 0, 1.2)
            rs = ps.add_run(subtitle)
            style_run(rs, FONT_BODY, size=self.font_role, italic=True, color=GREY)

    def bullet(self, segments):
        p = self.justify(self.doc.add_paragraph(style='List Bullet'))
        self.set_space(p, 0.35, 0.35)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        for text, bold in segments:
            self.append_rich_text(p, text, self.font_body, bold=bold, color=BLACK)

    def edu(self, inst, dates, detail):
        self.dated_line(inst, dates, title_font=self.font_edu, date_bold=True)
        if not detail:
            return
        pd = self.justify(self.doc.add_paragraph())
        self.set_space(pd, 0, 2)
        self.append_rich_text(
            pd, detail, self.font_edu_detail, bold=False, color=GREY
        )
        for run in pd.runs:
            run.italic = True

    def subhead(self, text):
        p = self.doc.add_paragraph()
        self.set_space(p, 6, 2)
        r = p.add_run(text)
        style_run(r, FONT_BODY, size=self.font_subhead, bold=True, color=ACCENT)

    def para(self, text, after=1.5):
        p = self.justify(self.doc.add_paragraph())
        self.set_space(p, 0, after)
        self.append_rich_text(p, text, self.font_body, bold=False, color=BLACK)

    def contact_details_line(self, parts, linkedin=None):
        """Single centered contact line: text parts and optional LinkedIn hyperlink."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_space(p, 0, 5)
        sep = '  |  '
        for idx, part in enumerate(parts):
            if idx:
                r = p.add_run(sep)
                style_run(r, FONT_BODY, size=self.font_contact, color=GREY)
            if '@' in part:
                self.append_hyperlink(
                    p,
                    part,
                    f'mailto:{part}',
                    self.font_contact,
                    bold=False,
                    color=ACCENT,
                )
            else:
                r = p.add_run(part)
                style_run(r, FONT_BODY, size=self.font_contact, color=GREY)
        if linkedin:
            display, url = linkedin
            if parts:
                r = p.add_run(sep)
                style_run(r, FONT_BODY, size=self.font_contact, color=GREY)
            self.append_hyperlink(
                p, display, url, self.font_contact, bold=False, color=ACCENT
            )

    def build_from_data(self, data, phone):
        self.name_heading(data['name'])
        if data.get('tagline'):
            self.tagline_line(data['tagline'])

        contact = data['contact_bits'][0] if data['contact_bits'] else ''
        linkedin = data.get('linkedin')
        if not linkedin:
            linkedin = extract_markdown_link(contact)
        contact = re.sub(
            r'\[([^\]]+)\]\([^)]+\)',
            '',
            contact,
        )
        contact = re.sub(r'\s*\|\s*$', '', contact).strip()
        contact = re.sub(r'\s*\|\s*\|', ' | ', contact)

        if '|' in contact:
            parts = [part.strip() for part in contact.split('|') if part.strip()]
        elif contact:
            parts = [contact]
        else:
            parts = []

        has_phone = any(
            part.replace('+', '').replace(' ', '').isdigit()
            or part.startswith('+')
            for part in parts
        )
        if not has_phone and phone:
            # Prefer email first, then phone, then remaining parts.
            if parts and '@' in parts[0]:
                parts = [parts[0], phone] + parts[1:]
            else:
                parts = [phone] + parts
        if not parts:
            parts = [phone]

        # Drop plain linkedin text leftovers; hyperlink is rendered separately.
        parts = [
            part for part in parts
            if 'linkedin.com' not in part.lower()
        ]
        self.contact_details_line(parts, linkedin=linkedin)

        has_freelance = any(
            section['title'].upper() == FREELANCE_SECTION.upper()
            for section in data['sections']
        )
        for section in data['sections']:
            title_upper = section['title'].upper()
            page_break = title_upper == FREELANCE_SECTION.upper()
            self.section(section['title'], page_break_before=page_break)

            if title_upper == 'PROFESSIONAL SUMMARY':
                paras = [b for b in section['blocks'] if b['type'] == 'para']
                for idx, block in enumerate(paras):
                    after = 0 if idx == len(paras) - 1 else 5
                    self.para(block['text'], after=after)
                continue

            if title_upper in SKILLS_SECTIONS:
                for block in section['blocks']:
                    if block['type'] == 'skill':
                        self.skill_line(block['label'], block['rest'])
                    elif block['type'] == 'bullet':
                        segments = block['segments']
                        if segments and segments[0][1] and segments[0][0].endswith(': '):
                            label = segments[0][0][:-2]
                            rest = ''.join(text for text, _ in segments[1:])
                            self.skill_line(label, rest)
                        else:
                            self.bullet(segments)
                continue

            if title_upper in {'PROFESSIONAL EXPERIENCE', FREELANCE_SECTION.upper()}:
                for block in section['blocks']:
                    if block['type'] != 'job':
                        continue
                    # Without freelance content, open page 2 at Symphony Teleca.
                    page_break_job = (
                        not has_freelance
                        and title_upper == 'PROFESSIONAL EXPERIENCE'
                        and block['company'].upper().startswith('SYMPHONY')
                    )
                    self.job(
                        block['company'],
                        block['role'],
                        block['dates'],
                        subtitle=block.get('subtitle'),
                        page_break_before=page_break_job,
                    )
                    for segments in block['bullets']:
                        self.bullet(segments)
                continue

            if title_upper in EDUCATION_SECTIONS:
                for block in section['blocks']:
                    if block['type'] == 'job':
                        self.edu(block['company'], block['dates'], block['role'])
                    elif block['type'] == 'subhead':
                        self.subhead(block['text'])
                    elif block['type'] == 'bullet':
                        self.bullet(block['segments'])
                continue

            if title_upper == 'LANGUAGES':
                for block in section['blocks']:
                    if block['type'] == 'bullet':
                        self.bullet(block['segments'])
                    elif block['type'] == 'para':
                        self.para(block['text'], after=0)


def convert_docx_to_pdf(docx_path, pdf_path):
    outdir = os.path.dirname(pdf_path)
    subprocess.run(
        [SOFFICE, '--headless', '--convert-to', 'pdf', '--outdir', outdir, docx_path],
        check=True,
        capture_output=True,
    )
    generated = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
    if generated != pdf_path:
        shutil.move(generated, pdf_path)


def page_text(pdf_path, page_num):
    result = subprocess.run(
        ['pdftotext', '-f', str(page_num), '-l', str(page_num), pdf_path, '-'],
        capture_output=True,
        text=True,
        check=True,
    )
    return result.stdout


def pdf_page_count(pdf_path):
    result = subprocess.run(
        ['pdfinfo', pdf_path],
        capture_output=True,
        text=True,
        check=True,
    )
    for line in result.stdout.splitlines():
        if line.startswith('Pages:'):
            return int(line.split(':', 1)[1].strip())
    raise RuntimeError(f'Could not determine page count for {pdf_path}')


def section_starts_on_page_2(pdf_path, marker):
    if pdf_page_count(pdf_path) != 2:
        return False
    page1 = page_text(pdf_path, 1).lower()
    page2 = page_text(pdf_path, 2).lower()
    if marker in page1:
        return False
    return marker in page2


def freelance_starts_on_page_2(pdf_path):
    return section_starts_on_page_2(pdf_path, FREELANCE_MARKER)


def layout_ok(pdf_path, exclude_freelance=False):
    if exclude_freelance:
        # Page 1: header through Netcracker. Page 2 opens at Symphony Teleca.
        if pdf_page_count(pdf_path) != 2:
            return False
        if not section_starts_on_page_2(pdf_path, 'symphony'):
            return False
        page2 = page_text(pdf_path, 2).lower()
        if 'netcracker' in page2:
            return False
        return True
    return freelance_starts_on_page_2(pdf_path)


def build_trial(data, phone, tmp_dir, tag, scale, line_spacing):
    docx_path = os.path.join(tmp_dir, f'cv-{tag}.docx')
    pdf_path = os.path.join(tmp_dir, f'cv-{tag}.pdf')
    builder = CvBuilder(scale, line_spacing=line_spacing)
    builder.build_from_data(data, phone)
    builder.doc.save(docx_path)
    convert_docx_to_pdf(docx_path, pdf_path)
    return pdf_path


def find_best_layout(data, phone, tmp_dir, exclude_freelance=False):
    scale_hi = 1.35 if exclude_freelance else 1.15
    lo, hi = 0.84, scale_hi
    best_scale = lo
    found = False
    for i in range(14):
        mid = (lo + hi) / 2
        pdf_path = build_trial(data, phone, tmp_dir, f'scale-{i}', mid, LINE_SPACING)
        if layout_ok(pdf_path, exclude_freelance=exclude_freelance):
            best_scale = mid
            lo = mid
            found = True
        else:
            hi = mid

    if not found:
        # Fall back to densest scale that still keeps a usable 2-page layout.
        for scale in (0.84, 0.88, 0.92, 0.96, 1.0):
            pdf_path = build_trial(data, phone, tmp_dir, f'fb-{scale}', scale, 0.96)
            if pdf_page_count(pdf_path) <= 2:
                return scale, 0.96
        return 0.84, 0.96

    lo, hi = (
        (max(1.0, LINE_SPACING), LINE_SPACING + 0.18)
        if exclude_freelance
        else (max(0.96, LINE_SPACING - 0.06), LINE_SPACING + 0.06)
    )
    best_line_spacing = LINE_SPACING
    for i in range(12):
        mid = (lo + hi) / 2
        pdf_path = build_trial(
            data, phone, tmp_dir, f'ls-{i}', best_scale, mid
        )
        if layout_ok(pdf_path, exclude_freelance=exclude_freelance):
            best_line_spacing = mid
            lo = mid
        else:
            hi = mid

    return best_scale, best_line_spacing


def generate_pdf(
    output_pdf,
    exclude_paragraphs=(),
    exclude_sections=(),
    remove_legacy=False,
    readme_name=README_NAME,
):
    secrets = load_secrets()
    if 'PHONE' not in secrets:
        raise SystemExit('secrets.txt must contain PHONE=...')
    phone = secrets['PHONE']

    readme_path = os.path.join(SCRIPT_DIR, readme_name)
    data = filter_cv_data(
        parse_readme(readme_path),
        exclude_paragraphs,
        exclude_sections,
    )
    exclude_freelance = not any(
        section['title'].upper() == FREELANCE_SECTION.upper()
        for section in data['sections']
    )

    out_pdf = os.path.join(SCRIPT_DIR, output_pdf)
    legacy = os.path.join(SCRIPT_DIR, 'beniamin-levin-cv.pdf')

    with tempfile.TemporaryDirectory() as tmp:
        best_scale, best_line_spacing = find_best_layout(
            data, phone, tmp, exclude_freelance=exclude_freelance
        )
        final_docx = os.path.join(tmp, 'cv-final.docx')
        final_pdf = os.path.join(tmp, 'cv-final.pdf')
        builder = CvBuilder(best_scale, line_spacing=best_line_spacing)
        builder.build_from_data(data, phone)
        builder.doc.save(final_docx)
        convert_docx_to_pdf(final_docx, final_pdf)

        if remove_legacy and os.path.lexists(legacy):
            os.remove(legacy)
        if os.path.lexists(out_pdf):
            os.remove(out_pdf)
        shutil.copy2(final_pdf, out_pdf)

    body_pt = BASE_FONT_BODY * best_scale
    print(
        f'saved {out_pdf} (scale={best_scale:.3f}, body={body_pt:.2f}pt, '
        f'line_spacing={best_line_spacing:.3f}, bottom_margin={MARGIN_BOTTOM.inches:.2f}in)'
    )


def main():
    parser = argparse.ArgumentParser(description='Generate CV PDF from README.md')
    parser.add_argument(
        '--output',
        default=DEFAULT_OUTPUT,
        help=f'Output PDF filename (default: {DEFAULT_OUTPUT})',
    )
    parser.add_argument(
        '--readme',
        default=README_NAME,
        help=f'Source markdown filename (default: {README_NAME})',
    )
    parser.add_argument(
        '--exclude-relocation',
        action='store_true',
        help='Omit the Netherlands relocation sentence from Professional Summary',
    )
    parser.add_argument(
        '--exclude-freelance',
        action='store_true',
        help='Omit the Freelance & Personal Projects section',
    )
    parser.add_argument(
        '--hands-on',
        action='store_true',
        help=f'Build hands-on CV from {HANDS_ON_README} → {HANDS_ON_OUTPUT}',
    )
    args = parser.parse_args()

    output = HANDS_ON_OUTPUT if args.hands_on else args.output
    readme = HANDS_ON_README if args.hands_on else args.readme
    exclude = (RELOCATION_LINE,) if args.exclude_relocation else ()
    exclude_sections = (FREELANCE_SECTION,) if args.exclude_freelance else ()
    generate_pdf(
        output,
        exclude_paragraphs=exclude,
        exclude_sections=exclude_sections,
        remove_legacy=output == DEFAULT_OUTPUT,
        readme_name=readme,
    )


if __name__ == '__main__':
    main()
